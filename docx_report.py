# -*- coding: utf-8 -*-
"""
동향 정리 워드 문서(.docx) 생성 — 메일 첨부용
서체: 조직 표준 Pretendard (한글 eastAsia 포함). 미보유 뷰어에서는 대체 폰트로 표시됨.
python-docx만 사용 → GitHub Actions에서 pip 설치로 동작.
"""

from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import config

FONT = "Pretendard"
NAVY = RGBColor(0x0C, 0x44, 0x7C)
ACCENT = RGBColor(0x18, 0x5F, 0xA5)
PRIMARY = RGBColor(0x1A, 0x1D, 0x21)
SECOND = RGBColor(0x4B, 0x51, 0x58)
MUTED = RGBColor(0x8B, 0x91, 0x99)
WD_KR = ["월", "화", "수", "목", "금", "토", "일"]
SHORT = {
    "산업 동향": "산업", "자사 동향": "자사",
    "경쟁사 동향": "경쟁사", "자회사 동향 (머큐리)": "자회사",
}


def _apply_font(run, name=FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT)


def _run(p, text, size=10.5, bold=False, color=PRIMARY):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    _apply_font(r)
    return r


def _shade(p, fill="F6F7F9"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _left_border(p, color="185FA5", size="18"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_hyperlink(p, url, text, size=9.0):
    part = p.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), FONT)
    rPr.append(rFonts)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "185FA5"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def _mmdd(pub):
    return f"{pub:%m.%d}" if pub else ""


def build_docx(display: dict, top: list, digests: dict, path: str, total: int) -> str:
    doc = Document()
    _set_default_font(doc)

    now = datetime.now(config.KST)

    # 제목
    p = doc.add_paragraph()
    _run(p, "아이즈비전 뉴스 클리핑", size=17, bold=True, color=NAVY)
    p = doc.add_paragraph()
    _run(p, f"{now:%Y년 %-m월 %-d일} ({WD_KR[now.weekday()]}) · 총 {total}건",
         size=10, color=MUTED)
    doc.add_paragraph()

    # Top 5
    p = doc.add_paragraph()
    _run(p, "오늘의 핵심 Top 5", size=13, bold=True, color=ACCENT)
    if top:
        for t in top:
            label = SHORT.get(t["cat_title"], t["cat_title"])
            pp = doc.add_paragraph()
            pp.paragraph_format.space_after = Pt(2)
            _run(pp, f"{t['rank']}. ", size=10.5, bold=True, color=ACCENT)
            _run(pp, t["headline"], size=10.5)
            _run(pp, f"  [{label}]", size=9, color=MUTED)
    else:
        _run(doc.add_paragraph(), "선별된 핵심 기사가 없습니다.", size=10, color=MUTED)
    doc.add_paragraph()

    # 섹션
    for cat in config.CATEGORIES:
        arts = display.get(cat["id"], [])

        # 섹션 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        _run(p, f"{cat['num']} {cat['title']}", size=13, bold=True, color=NAVY)
        _run(p, f"   {cat['subtitle']} · {len(arts)}건", size=9, color=MUTED)

        # 동향 요약 박스
        digest = (digests or {}).get(cat["id"], "")
        if digest:
            box = doc.add_paragraph()
            box.paragraph_format.space_before = Pt(4)
            box.paragraph_format.space_after = Pt(6)
            box.paragraph_format.left_indent = Pt(6)
            _shade(box)
            _left_border(box)
            _run(box, "이 섹션 한눈에  ", size=9, bold=True, color=ACCENT)
            _run(box, digest, size=9.5, color=SECOND)

        # 기사 목록
        if not arts:
            _run(doc.add_paragraph(), "해당 기간 수집된 기사가 없습니다.",
                 size=9.5, color=MUTED)
        for art in arts:
            pt = doc.add_paragraph()
            pt.paragraph_format.space_before = Pt(6)
            pt.paragraph_format.space_after = Pt(1)
            if art.get("badge"):
                _run(pt, f"[{art['badge']}] ", size=9, bold=True, color=ACCENT)
            _run(pt, art.get("title", ""), size=11, bold=True)

            ps = doc.add_paragraph()
            ps.paragraph_format.space_after = Pt(1)
            _run(ps, art.get("summary", ""), size=10, color=SECOND)

            src = doc.add_paragraph()
            src.paragraph_format.space_after = Pt(4)
            meta = " · ".join(x for x in [art.get("press", ""), _mmdd(art.get("pub"))] if x)
            _run(src, meta + "  ", size=9, color=MUTED)
            link = art.get("link")
            if link:
                _add_hyperlink(src, link, "원문보기")

        doc.add_paragraph()

    # 푸터
    fp = doc.add_paragraph()
    _run(fp, "본 리포트는 네이버 뉴스 기반 AI 자동 수집·요약이며, 정확한 내용은 원문 확인이 필요합니다.",
         size=8.5, color=MUTED)

    doc.save(path)
    return path
